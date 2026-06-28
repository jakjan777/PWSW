from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jakja\source\repos\PWSWlab10\docs\Exercise5_1_DFD_STRIDE.docx"


def set_cell_shading(cell, color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_dfd_table(doc):
    add_heading(doc, "Diagram przeplywu danych (DFD)", 2)

    doc.add_paragraph(
        "Aplikacja audytowa SecurityScanner (modul 1) w architekturze "
        "WinUI 3 + REST API + SQLite + Foundry Local + Azure SQL DB."
    )

    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ("Warstwa", "Komponent", "Opis / przeplyw"),
        ("Wejscie", "Uzytkownik", "Parametry skanu, raporty, konfiguracja"),
        ("Granica G1", "WinUI 3 Desktop App", "Dane wejsciowe moga byc zlosliwe"),
        ("Granica G2", "REST API (ASP.NET Core)", "Komunikacja sieciowa HTTPS"),
        ("Magazyn chmury", "Azure SQL DB", "Trwale dane audytu (G3)"),
        ("Magazyn lokalny", "SQLite", "Cache lokalny, historia skanow (G5)"),
        ("AI", "Foundry Local", "Analiza modelu bezpieczenstwa (G4)"),
    ]

    for r, row in enumerate(rows_data):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            if r == 0:
                set_cell_shading(cell, "D9E2F3")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Przeplyw danych: ").bold = True
    p.add_run(
        "Uzytkownik -> WinUI 3 -> REST API -> Azure SQL DB; "
        "WinUI 3 -> SQLite; WinUI 3 -> Foundry Local."
    )

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Powiazanie z modul 1: ").bold = True
    p2.add_run(
        "SubsystemMonitor.CheckAsync, Task.WhenAll, Task.WhenAny, "
        "SemaphoreSlim, FullScanRunner.RunFullScanAsync."
    )

    flow = doc.add_table(rows=5, cols=1)
    flow.style = "Table Grid"
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow_text = [
        "[ Uzytkownik ]",
        "        |",
        "        v  (Granica 1)",
        "[ WinUI 3 Desktop App ] -----> [ REST API ]",
        "    |              |                    |",
        "    v              v (Granica 2)        v (Granica 3)",
        "[ SQLite ]   [ Foundry Local ]   [ Azure SQL DB ]",
    ]
    merged = doc.add_table(rows=1, cols=1)
    merged.style = "Table Grid"
    cell = merged.rows[0].cells[0]
    cell.text = "\n".join([
        "[ Uzytkownik ]",
        "        |",
        "        v  (Granica 1: dane wejsciowe moga byc zlosliwe)",
        "[ WinUI 3 Desktop App ] -----> [ REST API (ASP.NET Core) ]",
        "    |              |                         |",
        "    v              v (Granica 2)             v (Granica 3)",
        "[ SQLite ]   [ Foundry Local ]        [ Azure SQL DB ]",
        "(G5)         (G4)                      (G3)",
    ])
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)

    doc.add_paragraph()


def add_trust_boundaries_table(doc):
    add_heading(doc, "Granice zaufania", 2)

    headers = ["Granica", "Obszar", "Glowne zagrozenia"]
    rows = [
        ["G1", "Uzytkownik <-> WinUI 3", "Zlosliwe dane wejsciowe, spoofing, XSS"],
        ["G2", "WinUI 3 <-> REST API", "Podszywanie sie, przechwycenie tokenu, MITM"],
        ["G3", "REST API <-> Azure SQL DB", "SQL injection, wyciek danych, eskalacja uprawnien"],
        ["G4", "WinUI 3 <-> Foundry Local", "Prompt injection, wyciek danych, OOM"],
        ["G5", "WinUI 3 <-> SQLite", "Manipulacja pliku bazy, brak szyfrowania at-rest"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h
        set_cell_shading(cell, "D9E2F3")

    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()


def add_stride_table(doc):
    add_heading(doc, "Tabela zagrozen STRIDE", 2)

    headers = ["#", "STRIDE", "Komponent", "Zagrozenie", "Mitygacja", "Ryzyko"]
    rows = [
        ["1", "S", "REST API", "Podswycie pod uzytkownika API", "OAuth 2.0, JWT", "Wysokie"],
        ["2", "T", "WinUI <-> REST API", "Modyfikacja danych w transmisji", "TLS 1.3, HTTPS", "Wysokie"],
        ["3", "T", "SQLite (lokalna)", "Manipulacja SQLite na dysku", "DPAPI, szyfrowanie at-rest", "Srednie"],
        ["4", "R", "REST API / WinUI", "Brak sladu operacji audytu", "Audit log, timestamp", "Srednie"],
        ["5", "I", "SQLite (lokalna)", "Wyciek danych z lokalnej bazy", "Szyfrowanie at-rest", "Srednie"],
        ["6", "I", "Foundry Local", "Model AI ujawnia dane treningowe", "Filtrowanie output", "Niskie"],
        ["7", "D", "REST API", "Przeciazenie REST API", "Rate limiting, WAF", "Wysokie"],
        ["8", "D", "Foundry Local", "OOM przez duze zapytanie AI", "Timeout, limit tokenow", "Srednie"],
        ["9", "E", "WinUI 3 (MSIX)", "Eskalacja przez blad w aplikacji", "AppContainer (MSIX)", "Wysokie"],
        ["10", "E", "REST API / Azure SQL", "SQL injection w wyszukiwarce", "Parameterized queries", "Wysokie"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(1), Cm(1.5), Cm(3), Cm(4), Cm(4), Cm(2)]
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h
        set_cell_shading(cell, "D9E2F3")

    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val
            if val == "Wysokie":
                set_cell_shading(table.rows[r].cells[5], "FCE4D6")
            elif val == "Srednie":
                set_cell_shading(table.rows[r].cells[5], "FFF2CC")
            elif val == "Niskie":
                set_cell_shading(table.rows[r].cells[5], "E2EFDA")

    doc.add_paragraph()


def add_stride_map_table(doc):
    add_heading(doc, "Mapa komponentow do kategorii STRIDE", 2)

    headers = ["Kategoria STRIDE", "Komponenty / granice"]
    rows = [
        ["Spoofing (S)", "REST API (G2)"],
        ["Tampering (T)", "Transmisja sieciowa (G2), SQLite na dysku (G5)"],
        ["Repudiation (R)", "Brak audit log (G1, G2)"],
        ["Information Disclosure (I)", "SQLite (G5), Foundry Local (G4)"],
        ["Denial of Service (D)", "REST API (G2), Foundry Local (G4)"],
        ["Elevation of Privilege (E)", "WinUI MSIX (G1), SQL injection (G3)"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"

    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h
        set_cell_shading(cell, "D9E2F3")

    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()


def add_priority_table(doc):
    add_heading(doc, "Priorytety mitygacji", 2)

    headers = ["Priorytet", "Zakres", "Dzialania"]
    rows = [
        ["Natychmiast (Wysokie)", "Przed dalszym rozwojem", "OAuth/JWT, HTTPS/TLS, rate limiting, parameterized SQL, AppContainer"],
        ["Przed release (Srednie)", "Przed wdrozeniem", "DPAPI, audit log, szyfrowanie SQLite, timeout AI"],
        ["Monitorowanie (Niskie)", "Ciagle", "Filtrowanie output modelu AI"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h
        set_cell_shading(cell, "D9E2F3")

    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val


def main():
    doc = Document()

    title = doc.add_heading("LAB-10  CWICZENIE 5.1", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Diagram DFD i analiza STRIDE dla aplikacji audytowej SecurityScanner"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_dfd_table(doc)
    add_trust_boundaries_table(doc)
    add_stride_table(doc)
    add_stride_map_table(doc)
    add_priority_table(doc)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
